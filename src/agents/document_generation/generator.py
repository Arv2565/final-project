import os
from typing import Dict, Any, Union
from docx import Document
from src.agents.document_generation.templates import DocumentTemplate

class DocumentGenerator:
    """
    Handles the generation of documents from templates.
    """

    def generate(self, template: DocumentTemplate, data: Dict[str, Any], output_path: str) -> str:
        """
        Generates a document by filling the template with data.
        
        Args:
            template: DocumentTemplate object.
            data: Dictionary of data to fill into placeholders.
            output_path: Path to save the generated document.
            
        Returns:
            Absolute path to the saved document.
        """
        # Resolve absolute path for template
        # Assuming run from project root
        if os.path.isabs(template.file_path):
             template_abs_path = template.file_path
        else:
             template_abs_path = os.path.abspath(template.file_path)

        if not os.path.exists(template_abs_path):
            raise FileNotFoundError(f"Template file not found at: {template_abs_path}")

        doc = Document(template_abs_path)
        
        # Merge default values with provided data
        filled_data = template.default_values.copy()
        filled_data.update(data)

        # Process paragraphs
        for paragraph in doc.paragraphs:
            self._replace_text_in_paragraph(paragraph, filled_data)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_text_in_paragraph(paragraph, filled_data)
        
        # Save
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        return os.path.abspath(output_path)

    def _replace_text_in_paragraph(self, paragraph, data: Dict[str, Any]):
        """
        Replaces placeholders in a paragraph. 
        Simple string replacement for now. 
        Note: python-docx runs can split text, simple replace might strictly work if placeholder is in one run.
        For robust replacement one might need to iterate runs or full text, but let's try simple replacement first.
        """
        # A simple approach that handles placeholders even if split across runs is tricky in docx.
        # However, often clear placeholders like {{KEY}} are distinct.
        # Let's try to replace in the full paragraph text if we can assign it back, 
        # but assigning to paragraph.text destroys formatting.
        # So we iterate runs. Ideally the template should have placeholders in single runs.
        
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}" # Assuming {{KEY}} format in docx
            # Check if placeholder exists in the whole paragraph text to decide if we need to search
            if placeholder in paragraph.text:
                # Naive run replacement (works if placeholder is in one run)
                replaced_in_run = False
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
                        replaced_in_run = True
                
                # If not found in single runs (split), we fall back to a bruteforce text replacement
                # This WILL clear formatting for that paragraph, which is a tradeoff.
                if not replaced_in_run:
                     paragraph.text = paragraph.text.replace(placeholder, str(value))

        # Also support non-bracketed keys if user prefers, but brackets are safer.
        # For this implementation I will assume the keys in `data` are the exact strings to find if they don't have brackets.
        # But commonly we map KEY -> {{KEY}}.
        # Let's adhere to the decision: the template registry lists keys like "COURT_NAME".
        # The docx SHOULD have {{COURT_NAME}}.
