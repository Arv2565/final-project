import re
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font_style(run, font_name="Helvetica", font_size=11, bold=False):
    """Apply generic font settings to a run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    # To strictly ensure font name in some Word versions
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)

def clean_text(text: str) -> str:
    """Clean text by replacing newline chars"""
    # Unlike PDF, python-docx handles newlines in paragraphs naturally if we split them, 
    # or we can keep them within a paragraph. 
    # For matching PDF generator logic, we process line by line.
    return text.replace('₹', 'Rs.')

def process_inline_formatting(paragraph, text: str):
    """
    Process bold markers (**text**) and ALL CAPS phrases.
    Appends runs to the given paragraph.
    """
    # This is a simplified parser. It handles **bold** and ALL CAPS.
    # It does NOT handle nested or complex overlapping.
    
    # Strategy: Tokenize by bold markers first.
    # re.split captures the separators if in parens.
    # r'(\*\*.*?\*\*)' splits into ['text', '**bold**', 'text']
    
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if not part:
            continue
            
        if part.startswith('**') and part.endswith('**'):
            # Bold content
            content = part[2:-2]
            run = paragraph.add_run(content)
            set_font_style(run, bold=True)
        else:
            # Check for ALL CAPS words (2+ chars)
            # We need to split this part again by all caps words
            # r'\b([A-Z]{2,}(?:\s+[A-Z]{2,})*)\b'
            
            sub_parts = re.split(r'\b([A-Z]{2,}(?:\s+[A-Z]{2,})*)\b', part)
            for sub_part in sub_parts:
                if not sub_part:
                    continue
                
                # Check if it matches the ALL CAPS pattern
                if re.match(r'^[A-Z]{2,}(?:\s+[A-Z]{2,})*$', sub_part):
                     # Bold it
                    run = paragraph.add_run(sub_part)
                    set_font_style(run, bold=True)
                else:
                    # Normal text
                    run = paragraph.add_run(sub_part)
                    set_font_style(run, bold=False)

def generate_docx_report(content: str, filename: str) -> str:
    """
    Generate a DOCX report from the given text content.
    
    Args:
        content: The text content to include (supports basic Markdown).
        filename: The output filename (including path).
        
    Returns:
        The absolute path to the generated DOCX.
    """
    output_path = Path(filename).resolve()
    # Change extension if needed, but usually passed correctly
    if output_path.suffix.lower() == '.pdf':
        output_path = output_path.with_suffix('.docx')
        
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    
    # ---------------------------------------------------------
    # 1. Page Setup (A4, 1 inch margins)
    # ---------------------------------------------------------
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ---------------------------------------------------------
    # 2. Style Definitions
    # ---------------------------------------------------------
    # We'll modify existing styles or create custom ones to match PDF logic.
    
    # Normal Style
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Helvetica' # Proxy for Aptos
    font.size = Pt(11)
    
    p_format = style_normal.paragraph_format
    p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format.space_after = Pt(6)
    p_format.line_spacing = 1.0 # Single spacing (approx)

    # Heading 1 (e.g., # Title)
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Helvetica-Bold'
    style_h1.font.size = Pt(11)
    style_h1.font.bold = True
    style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_h1.paragraph_format.space_after = Pt(12)
    style_h1.font.color.rgb = None # Default color (black usually)

    # Heading 2 (e.g., ## Subtitle)
    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = 'Helvetica-Bold'
    style_h2.font.size = Pt(11)
    style_h2.font.bold = True
    style_h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_h2.paragraph_format.space_after = Pt(12)
    style_h2.paragraph_format.space_before = Pt(0)
    style_h2.font.color.rgb = None

    # Title Style (Manual) for "Legal Analysis Report"
    # We can just use H1 or paragraphs with bold.

    # ---------------------------------------------------------
    # 3. Content Parsing
    # ---------------------------------------------------------
    
    # Add Title
    title_p = doc.add_paragraph("Legal Analysis Report")
    title_p.style = style_h1 # Reusing H1 style as base
    # Ensure it uses the font settings explicitly if style doesn't stick
    for run in title_p.runs:
        set_font_style(run, "Helvetica-Bold", 11, True)
    
    # Process Line by Line
    lines = content.split('\n')
    
    # List State
    current_indent_level = 0
    
    for line in lines:
        line = clean_text(line).strip() # Using strip() removes indentation for detection? 
        # Wait, PDF generator used regex and len checks on raw line. 
        # Let's keep raw indentation detection.
    
    lines = content.split('\n')
    for raw_line in lines:
        stripped_line = raw_line.strip()
        cleaned_line = clean_text(stripped_line)
        
        if not cleaned_line:
            # Blank line -> we can add empty paragraph or rely on space_after
            # To match PDF generator which added Spacer(1, 6), we might add small break
            # doc.add_paragraph("")
            continue
            
        # Calculate Indentation
        leading_spaces = len(raw_line) - len(raw_line.lstrip())
        indent_level = 0
        if leading_spaces >= 2: indent_level = 1
        if leading_spaces >= 6: indent_level = 2
        
        # Headers
        if cleaned_line.startswith('# '):
            p = doc.add_paragraph()
            p.style = style_h1
            text_content = cleaned_line[2:]
            process_inline_formatting(p, text_content)
            continue
            
        if cleaned_line.startswith('## '):
            p = doc.add_paragraph()
            p.style = style_h2
            text_content = cleaned_line[3:]
            process_inline_formatting(p, text_content)
            continue
            
        if cleaned_line.startswith('DRAFT OF '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            process_inline_formatting(p, cleaned_line)
            # Ensure whole paragraph is bold since process doesn't autocatch whole line unless regex matches
            # So let's force bold runs if not matched
            if not p.runs: # if process didn't add anything (unlikely)
                 pass
            return_runs = p.runs
            for r in return_runs:
                r.bold = True
                r.font.name = "Helvetica-Bold"
            continue

        # Lists
        is_list = False
        list_style = None # 'List Bullet' or 'List Number'
        content_text = cleaned_line
        
        # Bullets
        if cleaned_line.startswith('* ') or cleaned_line.startswith('- ') or cleaned_line.startswith('• '):
            is_list = True
            list_style = 'List Bullet'
            content_text = cleaned_line[2:].strip()
            
        # Numbered
        elif re.match(r'^(\d+)\.\s', cleaned_line):
            is_list = True
            list_style = 'List Number'
            match = re.match(r'^(\d+)\.\s(.*)', cleaned_line)
            content_text = match.group(2).strip()
            
        # Alpha
        elif re.match(r'^([a-z])\.\s', cleaned_line):
           is_list = True
           list_style = 'List Number' # Word generic number list usually handles sequences, but we might just use text bullet if precise control needed.
           # For simplicity, stick to List Number or manual text.
           # Let's use List Number for now.
           match = re.match(r'^([a-z])\.\s(.*)', cleaned_line)
           content_text = match.group(2).strip()

        if is_list:
            p = doc.add_paragraph()
            p.style = list_style
            # Indentation adjustment
            # Default List Indent in Word is usually 0.25 or 0.5 inch.
            # We want to mimic 24pt base + 24pt level.
            # 24pt = 1/3 inch.
            
            p.paragraph_format.left_indent = Inches(0.25 + (indent_level * 0.25))
            process_inline_formatting(p, content_text)
        else:
            # Normal Paragraph
            p = doc.add_paragraph()
            p.style = style_normal
            # Indentation for normal text if needed? Usually 0.
            if indent_level > 0:
                 p.paragraph_format.left_indent = Inches(indent_level * 0.25)
            process_inline_formatting(p, cleaned_line)
            
    # Footer Disclaimer
    doc.add_paragraph() # Spacer
    disclaimer = doc.add_paragraph("Disclaimer: This document is generated by an AI assistant. It is for informational purposes only and does not constitute professional legal advice.")
    disclaimer.style = 'Normal'
    for run in disclaimer.runs:
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = 0x808080 # RGB int? No check docx color
        # docx.shared.RGBColor(128, 128, 128)
    # Simplify color setting
    from docx.shared import RGBColor
    disclaimer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.save(str(output_path))
    return str(output_path)
